"""Export overview report as DOCX, PDF, or ZIP artifact bundle."""

import csv
import io
import re
import zipfile
from datetime import datetime


def _parse_markdown_lines(text):
    """Split markdown text into (style, content) pairs for simple rendering.

    Supports: headings (##), bold (**), bullet lists (- ), plain paragraphs.
    """
    lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            lines.append(("blank", ""))
        elif stripped.startswith("### "):
            lines.append(("h3", stripped[4:]))
        elif stripped.startswith("## "):
            lines.append(("h2", stripped[3:]))
        elif stripped.startswith("# "):
            lines.append(("h1", stripped[2:]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            lines.append(("bullet", stripped[2:]))
        else:
            lines.append(("para", stripped))
    return lines


def _render_markdown_docx(doc, description):
    """Render markdown description into a docx Document."""
    for style, content in _parse_markdown_lines(description):
        if style == "blank":
            doc.add_paragraph("")
        elif style == "h1":
            doc.add_heading(content, level=1)
        elif style == "h2":
            doc.add_heading(content, level=2)
        elif style == "h3":
            doc.add_heading(content, level=3)
        elif style == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", content)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)


def _set_page_landscape_wide(doc):
    """Set document to wide landscape with narrow margins for Google Docs pageless."""
    from docx.shared import Inches, Cm
    from docx.oxml.ns import qn

    section = doc.sections[0]
    # Landscape A4 wide
    section.page_width = Inches(16.5)
    section.page_height = Inches(11.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    return section


def _make_borderless_table(doc, rows, cols):
    """Create a borderless table for image grid layout."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    table = doc.add_table(rows=rows, cols=cols)
    tbl_pr = table._tbl.tblPr
    borders_elem = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_elem = OxmlElement(f"w:{edge}")
        edge_elem.set(qn("w:val"), "none")
        edge_elem.set(qn("w:sz"), "0")
        borders_elem.append(edge_elem)
    tbl_pr.append(borders_elem)

    # Set table to full width
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")  # 100% in fifths of a percent
    tbl_pr.append(tbl_w)

    return table


def export_docx(title, description, row_data, scenario_names, chart_sections=None):
    """Generate a DOCX report optimized for Google Docs pageless mode.

    Layout: wide landscape, no page breaks, large images in 2-col grid,
    full-width table.
    chart_sections: list of (section_title, [(chart_title, png_bytes), ...])
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = _set_page_landscape_wide(doc)

    # Available content width
    content_width = section.page_width - section.left_margin - section.right_margin
    img_half_width = Inches(6.5)  # Each image in 2-col layout

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts.runs[0].font.size = Pt(9)
    ts.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")

    # Description
    if description and description.strip():
        _render_markdown_docx(doc, description)
        doc.add_paragraph("")

    # Charts — 2-column layout per section, large images
    if chart_sections:
        for section_title, charts in chart_sections:
            doc.add_heading(section_title, level=1)

            for i in range(0, len(charts), 2):
                grid = _make_borderless_table(doc, 1, 2)

                # Left image
                _, png_left = charts[i]
                cell_left = grid.rows[0].cells[0]
                p_left = cell_left.paragraphs[0]
                p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_left = p_left.add_run()
                run_left.add_picture(io.BytesIO(png_left), width=img_half_width)

                # Right image (if exists)
                if i + 1 < len(charts):
                    _, png_right = charts[i + 1]
                    cell_right = grid.rows[0].cells[1]
                    p_right = cell_right.paragraphs[0]
                    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_right = p_right.add_run()
                    run_right.add_picture(io.BytesIO(png_right), width=img_half_width)

                # Small spacing paragraph
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(2)
                sp.paragraph_format.space_after = Pt(2)

    # Overview Table — no page break, just a heading
    doc.add_heading("Metrics Overview", level=1)

    col_names = ["Metric"] + scenario_names
    table = doc.add_table(rows=1, cols=len(col_names))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table to full width
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    table._tbl.tblPr.append(tbl_w)

    for i, col in enumerate(col_names):
        cell = table.rows[0].cells[i]
        cell.text = col
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for row in row_data:
        cells = table.add_row().cells
        cells[0].text = row.get("metric_pretty", "")
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        for i, sname in enumerate(scenario_names, 1):
            val = row.get(sname, "")
            cells[i].text = str(val) if val is not None else ""
            if cells[i].paragraphs[0].runs:
                cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(title, description, row_data, scenario_names, chart_sections=None):
    """Generate a PDF report and return bytes.

    Layout: title -> description -> charts (2-col, grouped by section) -> table.
    chart_sections: list of (section_title, [(chart_title, png_bytes), ...])
    """
    from fpdf import FPDF

    pdf = FPDF(orientation="L", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, title, ln=True, align="C")

    # Timestamp
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # Description
    if description and description.strip():
        for style, content in _parse_markdown_lines(description):
            if style == "blank":
                pdf.ln(3)
            elif style in ("h1", "h2", "h3"):
                sizes = {"h1": 14, "h2": 12, "h3": 10}
                pdf.set_font("Helvetica", "B", sizes[style])
                pdf.cell(0, 7, content, ln=True)
            elif style == "bullet":
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(8)
                pdf.cell(0, 5, f"- {content}", ln=True)
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.multi_cell(0, 5, content)
        pdf.ln(4)

    # Charts — 2-column layout per section
    page_w = pdf.w - pdf.l_margin - pdf.r_margin
    img_w = (page_w - 10) / 2  # 10mm gap between columns

    if chart_sections:
        for section_title, charts in chart_sections:
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, section_title, ln=True)
            pdf.ln(2)

            for i in range(0, len(charts), 2):
                # Check if we need a new page (enough room for ~90mm image height)
                if pdf.get_y() > pdf.h - 100:
                    pdf.add_page()

                y_pos = pdf.get_y()

                # Left image
                _, png_left = charts[i]
                pdf.image(io.BytesIO(png_left), x=pdf.l_margin, y=y_pos, w=img_w)

                # Right image (if exists)
                if i + 1 < len(charts):
                    _, png_right = charts[i + 1]
                    pdf.image(io.BytesIO(png_right), x=pdf.l_margin + img_w + 10, y=y_pos, w=img_w)

                # Move below the images
                pdf.set_y(y_pos + img_w * 0.57 + 5)  # approximate aspect ratio

    # Table
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Metrics Overview", ln=True)
    pdf.ln(2)
    _pdf_table(pdf, row_data, scenario_names)

    return bytes(pdf.output())


def _pdf_table(pdf, row_data, scenario_names):
    """Render the metrics table into the PDF."""
    col_names = ["Metric"] + scenario_names
    num_cols = len(col_names)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    metric_col_w = min(80, page_width * 0.3)
    data_col_w = (page_width - metric_col_w) / max(num_cols - 1, 1)
    col_widths = [metric_col_w] + [data_col_w] * (num_cols - 1)

    def _draw_header():
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_fill_color(66, 133, 244)
        pdf.set_text_color(255, 255, 255)
        for i, col in enumerate(col_names):
            pdf.cell(col_widths[i], 6, col[:30], border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(0, 0, 0)

    _draw_header()

    for idx, row in enumerate(row_data):
        if pdf.get_y() > pdf.h - 20:
            pdf.add_page()
            _draw_header()

        if idx % 2 == 0:
            pdf.set_fill_color(245, 245, 245)
        else:
            pdf.set_fill_color(255, 255, 255)

        metric = row.get("metric_pretty", "")[:40]
        pdf.cell(col_widths[0], 5, metric, border=1, fill=True)
        for i, sname in enumerate(scenario_names):
            val = row.get(sname, "")
            pdf.cell(col_widths[i + 1], 5, str(val) if val is not None else "",
                     border=1, fill=True, align="R")
        pdf.ln()


def _sanitize_filename(name):
    """Convert a chart title to a safe filename."""
    return re.sub(r"[^\w\s-]", "", name).strip().replace(" ", "_").lower()


def _section_folder_name(section_title):
    """Convert a section title to a folder name."""
    return _sanitize_filename(section_title)


def export_zip(row_data, scenario_names, chart_sections=None):
    """Generate a ZIP artifact bundle and return bytes.

    Structure:
        artifacts/
            ambient_platform_metrics/
                namespace_cpu_total.png
                namespace_memory_total.png
                ...
            cluster_metrics/
                cluster_cpu_usage_rate.png
                ...
            locust_load_test_metrics/
                avg_response_time.png
                ...
            overview_table.csv
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Charts organized by section folder
        if chart_sections:
            for section_title, charts in chart_sections:
                folder = _section_folder_name(section_title)
                for chart_title, png_bytes in charts:
                    filename = f"{_sanitize_filename(chart_title)}.png"
                    zf.writestr(f"artifacts/{folder}/{filename}", png_bytes)

        # Overview table as CSV
        csv_buf = io.StringIO()
        writer = csv.writer(csv_buf)
        writer.writerow(["Metric"] + scenario_names)
        for row in row_data:
            writer.writerow(
                [row.get("metric_pretty", "")]
                + [str(row.get(s, "")) if row.get(s) is not None else "" for s in scenario_names]
            )
        zf.writestr("artifacts/overview_table.csv", csv_buf.getvalue())

    return buf.getvalue()
